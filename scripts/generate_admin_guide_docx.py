# -*- coding: utf-8 -*-
"""Generate Vietnamese admin user guide as Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
    return h


def add_para(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("\n\n\nHƯỚNG DẪN SỬ DỤNG\nTRANG QUẢN TRỊ (ADMIN)\n")
    set_run_font(run, size=22, bold=True, color=RGBColor(0, 51, 102))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Website SVN Automation\n")
    set_run_font(run, size=14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("\nPhiên bản tài liệu: 1.0\nNgày cập nhật: Tháng 6/2025\n")
    set_run_font(run, size=11)

    doc.add_page_break()

    # Mục lục
    add_heading(doc, "MỤC LỤC", 1)
    toc_items = [
        "1. Giới thiệu",
        "2. Yêu cầu và truy cập hệ thống",
        "3. Đăng nhập và đăng xuất",
        "4. Giao diện chung của trang Admin",
        "5. Dashboard – Tổng quan",
        "6. Quản lý Sản phẩm & Giải pháp",
        "7. Quản lý Tin tức",
        "8. Quản lý Yêu cầu liên hệ",
        "9. Thư viện ảnh và Trình soạn thảo nội dung",
        "10. Lưu ý quan trọng khi vận hành",
        "11. Xử lý sự cố thường gặp",
    ]
    for item in toc_items:
        add_para(doc, item)
    doc.add_page_break()

    # 1. Giới thiệu
    add_heading(doc, "1. Giới thiệu", 1)
    add_para(
        doc,
        "Trang Quản trị (Admin Panel) của website SVN Automation là khu vực dành riêng cho "
        "quản trị viên để quản lý nội dung website công khai, bao gồm: sản phẩm, giải pháp, "
        "tin tức và theo dõi các yêu cầu liên hệ từ khách hàng.",
    )
    add_para(doc, "Các chức năng chính của trang Admin:", bold=True)
    add_bullet(doc, "Xem tổng quan thống kê nhanh trên Dashboard")
    add_bullet(doc, "Thêm, sửa, xóa sản phẩm và giải pháp")
    add_bullet(doc, "Viết, chỉnh sửa, xuất bản hoặc lưu nháp bài tin tức")
    add_bullet(doc, "Xem danh sách yêu cầu liên hệ từ form trên website")
    add_bullet(doc, "Upload và quản lý ảnh qua thư viện ảnh tích hợp")
    add_para(
        doc,
        "Lưu ý: Trang Admin không có liên kết trên website công khai. Bạn cần truy cập trực tiếp "
        "bằng địa chỉ URL (ví dụ: https://your-domain.com/admin/login).",
    )

    # 2. Yêu cầu và truy cập
    add_heading(doc, "2. Yêu cầu và truy cập hệ thống", 1)
    add_heading(doc, "2.1. Yêu cầu kỹ thuật", 2)
    add_bullet(doc, "Trình duyệt web hiện đại: Google Chrome, Microsoft Edge, Firefox hoặc Safari (phiên bản mới)")
    add_bullet(doc, "Kết nối Internet ổn định")
    add_bullet(doc, "Tài khoản quản trị do bộ phận IT hoặc quản trị hệ thống cấp")

    add_heading(doc, "2.2. Tài khoản Admin mặc định (lần cài đặt đầu tiên)", 2)
    add_para(
        doc,
        "Khi hệ thống được triển khai lần đầu, quản trị viên kỹ thuật sẽ chạy script tạo tài khoản. "
        "Tài khoản mặc định:",
    )
    add_table(
        doc,
        ["Thông tin", "Giá trị"],
        [
            ["Username", "admin"],
            ["Password", "svn@2025"],
        ],
    )
    add_para(
        doc,
        "Khuyến nghị: Đổi mật khẩu mặc định sau khi đăng nhập lần đầu (liên hệ bộ phận kỹ thuật nếu cần hỗ trợ).",
    )

    add_heading(doc, "2.3. Địa chỉ truy cập", 2)
    add_table(
        doc,
        ["Môi trường", "URL truy cập"],
        [
            ["Phát triển (local)", "http://localhost:3000/admin/login"],
            ["Production", "https://[tên-miền-của-bạn]/admin/login"],
        ],
    )

    # 3. Đăng nhập
    add_heading(doc, "3. Đăng nhập và đăng xuất", 1)
    add_heading(doc, "3.1. Các bước đăng nhập", 2)
    add_numbered(doc, "Mở trình duyệt và truy cập địa chỉ /admin/login")
    add_numbered(doc, "Nhập Username (tên đăng nhập) vào ô \"Username\"")
    add_numbered(doc, "Nhập Password (mật khẩu) vào ô \"Password\"")
    add_numbered(doc, "Nhấn nút \"Đăng nhập\" hoặc phím Enter trên bàn phím")
    add_numbered(doc, "Nếu thông tin đúng, hệ thống chuyển bạn đến trang Dashboard (/admin/dashboard)")
    add_para(doc, "Mẹo: Bấm biểu tượng con mắt bên cạnh ô mật khẩu để hiện/ẩn mật khẩu khi nhập.", indent=0.25)

    add_heading(doc, "3.2. Khi đăng nhập thất bại", 2)
    add_bullet(doc, "Thông báo \"Sai tài khoản hoặc mật khẩu\" – kiểm tra lại username và password (phân biệt chữ hoa/thường)")
    add_bullet(doc, "Kiểm tra phím Caps Lock có bật không")
    add_bullet(doc, "Liên hệ quản trị hệ thống nếu quên mật khẩu")

    add_heading(doc, "3.3. Đăng xuất", 2)
    add_para(doc, "Để đăng xuất an toàn:")
    add_numbered(doc, "Nhìn vào thanh menu bên trái (Sidebar)")
    add_numbered(doc, "Cuộn xuống dưới cùng")
    add_numbered(doc, "Nhấn nút \"Logout\"")
    add_para(doc, "Hệ thống sẽ xóa phiên đăng nhập và chuyển về trang đăng nhập.")

    add_heading(doc, "3.4. Thời hạn phiên đăng nhập", 2)
    add_para(
        doc,
        "Phiên đăng nhập có hiệu lực trong 24 giờ. Sau thời gian này, bạn cần đăng nhập lại. "
        "Nếu thấy thông báo \"Phiên đăng nhập hết hạn\", hãy đăng nhập lại.",
    )

    # 4. Giao diện chung
    add_heading(doc, "4. Giao diện chung của trang Admin", 1)
    add_para(doc, "Sau khi đăng nhập, giao diện Admin gồm ba phần chính:", bold=True)

    add_heading(doc, "4.1. Sidebar (Menu bên trái)", 2)
    add_table(
        doc,
        ["Mục menu", "Chức năng", "Đường dẫn"],
        [
            ["Dashboard", "Tổng quan thống kê", "/admin/dashboard"],
            ["Products", "Quản lý sản phẩm & giải pháp", "/admin/products"],
            ["News", "Quản lý tin tức", "/admin/news"],
            ["Liên hệ", "Xem yêu cầu liên hệ", "/admin/contacts"],
        ],
    )
    add_para(doc, "Phía dưới sidebar hiển thị thông tin tài khoản và nút Logout.")

    add_heading(doc, "4.2. Header (Thanh tiêu đề trên)", 2)
    add_para(doc, "Hiển thị tên trang hiện tại (Dashboard, Products, News, Liên hệ) và avatar quản trị viên.")

    add_heading(doc, "4.3. Khu vực nội dung chính", 2)
    add_para(doc, "Phần giữa màn hình hiển thị nội dung tương ứng với mục menu bạn chọn.")

    doc.add_page_break()

    # 5. Dashboard
    add_heading(doc, "5. Dashboard – Tổng quan", 1)
    add_para(
        doc,
        "Dashboard là trang đầu tiên sau khi đăng nhập, cung cấp cái nhìn nhanh về tình trạng website.",
    )

    add_heading(doc, "5.1. Các thẻ thống kê", 2)
    add_table(
        doc,
        ["Thẻ thống kê", "Ý nghĩa"],
        [
            ["Tổng sản phẩm", "Tổng số sản phẩm và giải pháp trong hệ thống"],
            ["Tổng bài viết", "Tổng số bài tin tức (cả nháp và đã xuất bản)"],
            ["Đã xuất bản", "Số bài tin tức đang hiển thị công khai"],
            ["Sản phẩm active", "Số sản phẩm/giải pháp đang bật hiển thị"],
            ["Yêu cầu liên hệ", "Tổng số form liên hệ khách hàng đã gửi"],
        ],
    )

    add_heading(doc, "5.2. Nút hành động nhanh", 2)
    add_bullet(doc, "+ Thêm sản phẩm – Chuyển đến trang quản lý sản phẩm và mở form tạo mới")
    add_bullet(doc, "+ Viết bài – Chuyển đến trang quản lý tin tức và mở form viết bài mới")

    add_heading(doc, "5.3. Bảng dữ liệu gần đây", 2)
    add_para(doc, "Dashboard hiển thị ba bảng:")
    add_bullet(doc, "5 sản phẩm mới nhất – cột: Tên, Ngày tạo, Trạng thái (Active/Inactive)")
    add_bullet(doc, "5 bài viết mới nhất – cột: Tên, Ngày tạo, Trạng thái (Published/Draft)")
    add_bullet(doc, "8 yêu cầu liên hệ gần đây – cột: Họ tên, Email, Công ty, Mục đích, Nguồn, Ngày gửi")
    add_para(doc, "Mỗi bảng có liên kết \"Xem tất cả\" để chuyển đến trang quản lý chi tiết.")

    doc.add_page_break()

    # 6. Products
    add_heading(doc, "6. Quản lý Sản phẩm & Giải pháp", 1)
    add_para(
        doc,
        "Trang Products cho phép quản lý toàn bộ sản phẩm (category: Sản phẩm) và giải pháp "
        "(category: Giải pháp) hiển thị trên website công khai.",
    )

    add_heading(doc, "6.1. Xem danh sách sản phẩm", 2)
    add_numbered(doc, "Từ sidebar, nhấn \"Products\" hoặc truy cập /admin/products")
    add_numbered(doc, "Bảng danh sách hiển thị các cột: STT, Ảnh, Tiêu đề, Danh mục, Trạng thái, Tạo lúc, Hành động")
    add_numbered(doc, "Nhấn nút \"Refresh\" để tải lại danh sách mới nhất")

    add_heading(doc, "6.2. Thêm sản phẩm mới", 2)
    add_numbered(doc, "Nhấn nút \"+ Thêm sản phẩm\" ở góc trên bên phải")
    add_numbered(doc, "Điền các trường trong form (xem bảng bên dưới)")
    add_numbered(doc, "Nhấn \"Lưu\" để lưu hoặc \"Quay lại\" để hủy")
    add_numbered(doc, "Sau khi lưu thành công, hệ thống quay về danh sách")

    add_table(
        doc,
        ["Trường", "Bắt buộc", "Mô tả"],
        [
            ["Tiêu đề", "Có (*)", "Tên sản phẩm hoặc giải pháp"],
            ["Mô tả ngắn (danh sách)", "Không", "Hiển thị trên trang danh sách và thẻ sản phẩm"],
            ["Tóm tắt (trang chi tiết)", "Không", "Đoạn dẫn ngắn dưới tiêu đề trang chi tiết"],
            ["Slug URL", "Tự động", "Chỉ hiển thị khi sửa; tự sinh từ tiêu đề, không chỉnh sửa được"],
            ["Nội dung chi tiết", "Không", "Nội dung đầy đủ, dùng trình soạn thảo rich text"],
            ["Danh mục", "Có", "Chọn \"Sản phẩm\" hoặc \"Giải pháp\""],
            ["Thứ tự", "Có", "Số thứ tự sắp xếp (số nhỏ hiển thị trước, mặc định: 0)"],
            ["Trạng thái", "Có", "Bật = Hiển thị trên web; Tắt = Ẩn khỏi website"],
            ["Tự động dịch tiếng Anh", "Không", "Mặc định BẬT – tự dịch tiêu đề và nội dung sang tiếng Anh khi lưu"],
            ["Ảnh sản phẩm", "Không", "Upload ảnh JPG/PNG/WEBP làm ảnh đại diện"],
        ],
    )

    add_heading(doc, "6.3. Sửa sản phẩm", 2)
    add_numbered(doc, "Trong danh sách, tìm sản phẩm cần sửa")
    add_numbered(doc, "Nhấn biểu tượng bút chì (✏️) ở cột Hành động")
    add_numbered(doc, "Chỉnh sửa các trường cần thiết")
    add_numbered(doc, "Nhấn \"Lưu\"")

    add_heading(doc, "6.4. Xóa sản phẩm", 2)
    add_numbered(doc, "Nhấn biểu tượng thùng rác (🗑️) ở cột Hành động")
    add_numbered(doc, "Xác nhận trong hộp thoại: \"Xóa sản phẩm \\\"...\\\"?\"")
    add_numbered(doc, "Nhấn OK/Xác nhận để xóa vĩnh viễn")
    add_para(doc, "Cảnh báo: Thao tác xóa không thể hoàn tác.", bold=True)

    add_heading(doc, "6.5. Ý nghĩa trạng thái và danh mục", 2)
    add_table(
        doc,
        ["Badge hiển thị", "Ý nghĩa"],
        [
            ["Hiển thị", "Sản phẩm đang active, khách truy cập thấy trên website"],
            ["Ẩn", "Sản phẩm bị ẩn, không hiển thị công khai"],
            ["Sản phẩm", "Thuộc danh mục sản phẩm"],
            ["Giải pháp", "Thuộc danh mục giải pháp"],
        ],
    )

    doc.add_page_break()

    # 7. News
    add_heading(doc, "7. Quản lý Tin tức", 1)
    add_para(doc, "Trang News dùng để tạo, chỉnh sửa, xuất bản hoặc lưu nháp các bài viết tin tức.")

    add_heading(doc, "7.1. Xem danh sách tin tức", 2)
    add_numbered(doc, "Từ sidebar, nhấn \"News\" hoặc truy cập /admin/news")
    add_numbered(doc, "Danh sách hiển thị 10 bài mỗi trang")
    add_numbered(doc, "Dùng tab lọc: \"Tất cả\" / \"Đã xuất bản\" / \"Bản nháp\"")
    add_numbered(doc, "Dùng nút \"← Trước\" và \"Sau →\" để chuyển trang")

    add_heading(doc, "7.2. Viết bài mới", 2)
    add_numbered(doc, "Nhấn \"+ Viết bài\"")
    add_numbered(doc, "Nhập Tiêu đề (bắt buộc) – slug tự sinh khi rời khỏi ô tiêu đề")
    add_numbered(doc, "Điền Tóm tắt, Nội dung (rich editor), Ngày xuất bản (nếu có)")
    add_numbered(doc, "Chọn Trạng thái: Xuất bản (hiển thị công khai) hoặc Bản nháp (chỉ lưu nội bộ)")
    add_numbered(doc, "Bật/tắt \"Tự động dịch tiếng Anh\" nếu cần")
    add_numbered(doc, "Upload ảnh đại diện bài viết")
    add_numbered(doc, "Nhấn \"Lưu\"")

    add_table(
        doc,
        ["Trường", "Bắt buộc", "Mô tả"],
        [
            ["Tiêu đề", "Có (*)", "Tiêu đề bài viết"],
            ["Slug", "Tự động", "Đường dẫn URL thân thiện, sinh từ tiêu đề"],
            ["Tóm tắt", "Không", "Mô tả ngắn hiển thị trên danh sách tin"],
            ["Nội dung", "Không", "Nội dung đầy đủ bài viết"],
            ["Ngày xuất bản", "Không", "Ngày hiển thị trên website (chọn từ lịch)"],
            ["Trạng thái", "Có", "Xuất bản hoặc Bản nháp"],
            ["Tự động dịch tiếng Anh", "Không", "Mặc định BẬT"],
            ["Ảnh bài viết", "Không", "Ảnh thumbnail cho bài viết"],
        ],
    )

    add_heading(doc, "7.3. Sửa và xóa bài viết", 2)
    add_para(doc, "Sửa bài:", bold=True)
    add_bullet(doc, "Nhấn biểu tượng bút chì trên dòng bài viết cần sửa")
    add_bullet(doc, "Chỉnh sửa và nhấn \"Lưu\"")
    add_para(doc, "Xóa bài:", bold=True)
    add_bullet(doc, "Nhấn biểu tượng thùng rác")
    add_bullet(doc, "Xác nhận: \"Xóa bài viết \\\"...\\\"?\"")

    doc.add_page_break()

    # 8. Contacts
    add_heading(doc, "8. Quản lý Yêu cầu liên hệ", 1)
    add_para(
        doc,
        "Trang Liên hệ hiển thị các form liên hệ khách hàng gửi từ website. "
        "Đây là trang CHỈ ĐỌC – không có chức năng sửa hoặc xóa yêu cầu.",
    )

    add_heading(doc, "8.1. Cách xem yêu cầu liên hệ", 2)
    add_numbered(doc, "Từ sidebar, nhấn \"Liên hệ\" hoặc truy cập /admin/contacts")
    add_numbered(doc, "Xem bảng danh sách: STT, Thời gian, Họ tên, Email, Công ty, Chức danh, Mục đích, Quốc gia, Nguồn")
    add_numbered(doc, "Nhấn nút mũi tên (▶) trên một dòng để mở rộng chi tiết")
    add_numbered(doc, "Xem thêm: Ngành/lĩnh vực và Nhu cầu/nội dung đầy đủ")
    add_numbered(doc, "Nhấn lại để thu gọn")
    add_numbered(doc, "Nhấn \"Làm mới\" để tải lại dữ liệu")

    add_heading(doc, "8.2. Phân trang", 2)
    add_para(doc, "Mỗi trang hiển thị 25 yêu cầu. Dùng nút phân trang để xem các trang tiếp theo.")

    add_heading(doc, "8.3. Nguồn yêu cầu", 2)
    add_table(
        doc,
        ["Badge hiển thị", "Nguồn gốc"],
        [
            ["CTA trang chủ", "Khách gửi từ nút kêu gọi hành động trên trang chủ"],
            ["Trang liên hệ", "Khách gửi từ trang Liên hệ"],
            ["Khác", "Nguồn khác hoặc không xác định"],
        ],
    )

    add_heading(doc, "8.4. Liên hệ qua email", 2)
    add_para(doc, "Nhấn vào địa chỉ Email trong bảng để mở ứng dụng email mặc định (mailto).")

    doc.add_page_break()

    # 9. Rich editor & image library
    add_heading(doc, "9. Thư viện ảnh và Trình soạn thảo nội dung", 1)
    add_para(
        doc,
        "Khi soạn Nội dung chi tiết (sản phẩm) hoặc Nội dung (tin tức), bạn dùng trình soạn thảo "
        "rich text tích hợp với nhiều công cụ định dạng.",
    )

    add_heading(doc, "9.1. Công cụ định dạng văn bản", 2)
    add_bullet(doc, "Tiêu đề (Heading 1, 2, 3...)")
    add_bullet(doc, "In đậm, in nghiêng, gạch chân")
    add_bullet(doc, "Danh sách có số và danh sách dấu chấm")
    add_bullet(doc, "Căn lề trái, giữa, phải")
    add_bullet(doc, "Chèn liên kết (link)")
    add_bullet(doc, "Đổi màu chữ")
    add_bullet(doc, "Trích dẫn (blockquote) và mã nguồn (code)")

    add_heading(doc, "9.2. Chèn ảnh vào nội dung", 2)
    add_numbered(doc, "Trong trình soạn thảo, nhấn biểu tượng ảnh trên thanh công cụ")
    add_numbered(doc, "Cửa sổ \"Thư viện ảnh\" mở ra")

    add_heading(doc, "9.3. Tab Thư viện", 3)
    add_bullet(doc, "Xem lưới các ảnh đã upload trước đó")
    add_bullet(doc, "\"Làm mới\" – tải lại danh sách ảnh")
    add_bullet(doc, "\"+ Tải ảnh\" – chuyển sang tab Upload")
    add_bullet(doc, "Click chọn một hoặc nhiều ảnh (ảnh được chọn sẽ được đánh dấu)")
    add_bullet(doc, "\"Chèn vào nội dung\" – chèn ảnh đã chọn vào vị trí con trỏ trong bài viết")
    add_bullet(doc, "\"Xóa đã chọn\" – xóa ảnh khỏi Cloudinary (có hộp thoại xác nhận)")

    add_heading(doc, "9.4. Tab Tải lên", 3)
    add_bullet(doc, "Chọn một hoặc nhiều file ảnh (JPG, PNG, WEBP)")
    add_bullet(doc, "Dung lượng tối đa: 5MB mỗi ảnh")
    add_bullet(doc, "Sau khi upload thành công, tự động quay về tab Thư viện")

    add_heading(doc, "9.5. Upload ảnh đại diện", 2)
    add_para(
        doc,
        "Ngoài ảnh trong nội dung, mỗi sản phẩm/tin tức có trường \"Ảnh sản phẩm\" hoặc "
        "\"Ảnh bài viết\" riêng – dùng làm ảnh thumbnail trên danh sách và thẻ hiển thị.",
    )

    doc.add_page_break()

    # 10. Lưu ý
    add_heading(doc, "10. Lưu ý quan trọng khi vận hành", 1)
    add_bullet(
        doc,
        "Cloudinary bắt buộc: Mọi ảnh upload (ảnh đại diện, ảnh trong nội dung) đều lưu trên Cloudinary. "
        "Nếu upload thất bại, liên hệ bộ phận kỹ thuật kiểm tra cấu hình CLOUDINARY_URL.",
    )
    add_bullet(
        doc,
        "Tự động dịch tiếng Anh: Khi bật, hệ thống tự dịch tiêu đề và nội dung sang tiếng Anh khi lưu. "
        "Tắt nếu bạn muốn tự nhập nội dung tiếng Anh.",
    )
    add_bullet(
        doc,
        "Sản phẩm và Giải pháp dùng chung trang Products – phân biệt bằng trường Danh mục.",
    )
    add_bullet(
        doc,
        "Bản nháp tin tức và sản phẩm ẩn không hiển thị trên website công khai cho khách truy cập.",
    )
    add_bullet(
        doc,
        "Không chia sẻ tài khoản admin với người không có thẩm quyền.",
    )
    add_bullet(
        doc,
        "Luôn đăng xuất khi rời khỏi máy tính dùng chung.",
    )

    # 11. Troubleshooting
    add_heading(doc, "11. Xử lý sự cố thường gặp", 1)
    add_table(
        doc,
        ["Vấn đề", "Nguyên nhân có thể", "Cách xử lý"],
        [
            [
                "Không đăng nhập được",
                "Sai username/password",
                "Kiểm tra lại thông tin, tắt Caps Lock",
            ],
            [
                "Phiên đăng nhập hết hạn",
                "Token JWT hết hạn sau 24h",
                "Đăng nhập lại",
            ],
            [
                "Upload ảnh thất bại",
                "Cloudinary chưa cấu hình hoặc file quá lớn",
                "Kiểm tra định dạng JPG/PNG/WEBP, dung lượng ≤ 5MB; liên hệ IT",
            ],
            [
                "Lưu sản phẩm/tin thất bại",
                "Thiếu trường bắt buộc hoặc lỗi mạng",
                "Điền đủ Tiêu đề, kiểm tra kết nối Internet, thử Refresh và lưu lại",
            ],
            [
                "Không thấy sản phẩm trên website",
                "Trạng thái đang \"Ẩn\"",
                "Vào Products Admin, sửa sản phẩm và bật trạng thái Hiển thị",
            ],
            [
                "Không thấy tin tức trên website",
                "Bài đang ở trạng thái Bản nháp",
                "Vào News Admin, sửa bài và chuyển sang Xuất bản",
            ],
        ],
    )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— Hết tài liệu —\nSVN Automation Admin Guide v1.0")
    set_run_font(run, size=10, color=RGBColor(128, 128, 128))

    return doc


def main():
    out_dir = Path(__file__).resolve().parent.parent / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "HUONG-DAN-SU-DUNG-TRANG-ADMIN.docx"
    doc = build_document()
    doc.save(str(out_path))
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()
